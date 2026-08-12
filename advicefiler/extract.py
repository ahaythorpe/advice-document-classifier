"""Step 1 of the four steps: document -> text.

The v0 harness skipped this entirely and read pre-extracted text out of a JSON
file, which quietly assumed every real document yields clean text. Most do. The
ones that do not are the ones that matter: SYSTEM.md section 7 names bad scans as
a failure mode precisely because a crooked photocopy produces *something*, and
something is enough for a keyword matcher to score on.

So extraction reports quality alongside text, and the pipeline is expected to
refuse to classify confidently when quality is poor. A document we cannot read is
not a document we may guess about.

PDF and Word support are optional imports. With no dependencies installed the
package still runs end to end against sample_documents.json.
"""

from __future__ import annotations

import hashlib
import os
import re
import unicodedata
from typing import Any, Dict, List, Optional, Tuple

# Optional backends. Absence is normal, not an error, until a real file arrives.
try:
    import pypdf  # type: ignore
except ImportError:
    pypdf = None

try:
    import pdfplumber  # type: ignore
except ImportError:
    pdfplumber = None

try:
    import docx  # type: ignore  # python-docx
except ImportError:
    docx = None


PDF_SUFFIXES = (".pdf",)
WORD_SUFFIXES = (".docx",)
TEXT_SUFFIXES = (".txt", ".text", ".md")

SUPPORTED_SUFFIXES = PDF_SUFFIXES + WORD_SUFFIXES + TEXT_SUFFIXES


class ExtractionError(RuntimeError):
    pass


class MissingBackend(ExtractionError):
    """A real document arrived and the library that reads it is not installed."""


# ---------------------------------------------------------------------------
# Quality
# ---------------------------------------------------------------------------

# Thresholds for calling a scan unreadable. These are deliberately generous:
# the cost of quarantining a readable document is one human glance, the cost of
# confidently classifying an unreadable one is a misfiled compliance record.
MIN_CHARS_TOTAL = 120
MIN_CHARS_PER_PAGE = 80
MIN_ALPHA_RATIO = 0.55
MIN_DICTIONARY_HIT_RATIO = 0.35
MAX_SINGLE_CHAR_TOKEN_RATIO = 0.35

# Words common enough that their absence from a page of "text" means the text is
# not English prose. Deliberately mundane and domain-free.
_COMMON_WORDS = frozenset("""
the of and to in for a is are be on with this that you your we our it as at by or
not from will may any has have been which if all can other more please client
""".split())


class Quality(object):
    """A verdict on whether extracted text is worth classifying."""

    def __init__(
        self,
        readable: bool,
        score: float,
        reasons: List[str],
        stats: Dict[str, Any],
    ) -> None:
        self.readable = readable
        self.score = score
        self.reasons = reasons
        self.stats = stats

    def __repr__(self) -> str:  # pragma: no cover - debugging aid
        return "<Quality readable=%s score=%.2f %s>" % (
            self.readable, self.score, ";".join(self.reasons) or "ok")


def assess_quality(text: str, page_count: int = 1) -> Quality:
    """Score extracted text for readability.

    Returns a score in [0, 1] and the reasons it lost points. The pipeline uses
    the score as a ceiling on classification confidence, so an unreadable
    document cannot be filed confidently even when it happens to contain a
    keyword that scores well.
    """
    reasons = []  # type: List[str]
    stripped = text.strip()
    total = len(stripped)
    page_count = max(1, page_count)

    if total == 0:
        return Quality(False, 0.0, ["no text extracted at all"],
                       {"chars": 0, "pages": page_count})

    letters = sum(1 for ch in stripped if ch.isalpha())
    alpha_ratio = float(letters) / total
    per_page = float(total) / page_count

    tokens = re.findall(r"[A-Za-z']+", stripped.lower())
    token_count = len(tokens)
    single_char = sum(1 for t in tokens if len(t) == 1)
    single_ratio = float(single_char) / token_count if token_count else 1.0
    dictionary_hits = sum(1 for t in tokens if t in _COMMON_WORDS)
    # Measured against a nominal 60-common-words-per-1000-tokens baseline.
    expected_hits = max(1.0, token_count * 0.04)
    dict_ratio = min(1.0, dictionary_hits / expected_hits)

    score = 1.0

    if total < MIN_CHARS_TOTAL:
        reasons.append("only %d characters extracted" % total)
        score *= max(0.05, float(total) / MIN_CHARS_TOTAL)
    if per_page < MIN_CHARS_PER_PAGE:
        reasons.append("%.0f characters per page (expected >= %d)"
                       % (per_page, MIN_CHARS_PER_PAGE))
        score *= max(0.05, per_page / MIN_CHARS_PER_PAGE)
    if alpha_ratio < MIN_ALPHA_RATIO:
        reasons.append("only %.0f%% of characters are letters" % (alpha_ratio * 100))
        score *= max(0.05, alpha_ratio / MIN_ALPHA_RATIO)
    if single_ratio > MAX_SINGLE_CHAR_TOKEN_RATIO:
        reasons.append("%.0f%% of words are single characters (OCR shrapnel)"
                       % (single_ratio * 100))
        score *= max(0.05, MAX_SINGLE_CHAR_TOKEN_RATIO / max(single_ratio, 0.01))
    if dict_ratio < MIN_DICTIONARY_HIT_RATIO:
        reasons.append("almost no common English words (%d in %d tokens)"
                       % (dictionary_hits, token_count))
        score *= max(0.05, dict_ratio / MIN_DICTIONARY_HIT_RATIO)

    score = max(0.0, min(1.0, score))
    stats = {
        "chars": total,
        "pages": page_count,
        "chars_per_page": round(per_page, 1),
        "alpha_ratio": round(alpha_ratio, 3),
        "tokens": token_count,
        "single_char_token_ratio": round(single_ratio, 3),
        "common_word_hits": dictionary_hits,
    }
    return Quality(score >= 0.5, round(score, 3), reasons, stats)


# ---------------------------------------------------------------------------
# Extracted document
# ---------------------------------------------------------------------------

class ExtractedDocument(object):
    """Text plus provenance, handed to the classifier."""

    def __init__(
        self,
        name: str,
        text: str,
        page_count: int = 1,
        source_path: Optional[str] = None,
        backend: str = "unknown",
        quality: Optional[Quality] = None,
    ) -> None:
        self.name = name
        self.text = text
        self.page_count = page_count
        self.source_path = source_path
        self.backend = backend
        self.quality = quality or assess_quality(text, page_count)

    @property
    def content_id(self) -> str:
        """Short stable id from the text, used for idempotent re-filing."""
        digest = hashlib.sha256(self.text.encode("utf-8")).hexdigest()
        return digest[:16]

    @property
    def suffix(self) -> str:
        if self.source_path:
            return os.path.splitext(self.source_path)[1]
        return os.path.splitext(self.name)[1]

    def __repr__(self) -> str:  # pragma: no cover - debugging aid
        return "<ExtractedDocument %s %d chars via %s>" % (
            self.name, len(self.text), self.backend)


def _normalise(text: str) -> str:
    """Tidy extractor output without destroying structure.

    Line breaks are load-bearing: the classifier's title-position rule looks for
    a title on a line of its own, so lines are preserved and only runs of
    horizontal whitespace are collapsed.
    """
    text = unicodedata.normalize("NFKC", text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Ligatures and typographic quotes that break literal pattern matching.
    for bad, good in (("ﬁ", "fi"), ("ﬂ", "fl"), ("’", "'"),
                      ("‘", "'"), ("“", '"'), ("”", '"'),
                      ("–", "-"), ("—", "-"), (" ", " ")):
        text = text.replace(bad, good)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# ---------------------------------------------------------------------------
# Backends
# ---------------------------------------------------------------------------

def _extract_pdf(path: str) -> Tuple[str, int, str]:
    """Try pypdf first, fall back to pdfplumber.

    pypdf is fast and fine for text-native advice records. pdfplumber is slower
    but recovers more from the tabular layouts fact finds are built from, so it
    is worth a second pass when the first yields almost nothing.
    """
    errors = []  # type: List[str]
    best = ("", 0, "none")

    if pypdf is not None:
        try:
            reader = pypdf.PdfReader(path)
            pages = [(p.extract_text() or "") for p in reader.pages]
            best = ("\n\n".join(pages), len(pages), "pypdf")
        except Exception as exc:  # noqa: BLE001 - backend failure is data, not a crash
            errors.append("pypdf: %s" % exc)

    needs_fallback = len(best[0].strip()) < MIN_CHARS_TOTAL
    if needs_fallback and pdfplumber is not None:
        try:
            with pdfplumber.open(path) as pdf:
                pages = [(p.extract_text() or "") for p in pdf.pages]
            candidate = ("\n\n".join(pages), len(pages), "pdfplumber")
            if len(candidate[0].strip()) > len(best[0].strip()):
                best = candidate
        except Exception as exc:  # noqa: BLE001
            errors.append("pdfplumber: %s" % exc)

    if best[2] == "none":
        if pypdf is None and pdfplumber is None:
            raise MissingBackend(
                "cannot read %s — no PDF backend installed. "
                "Run: python3 -m pip install -r requirements.txt" % os.path.basename(path))
        raise ExtractionError("PDF extraction failed for %s (%s)"
                              % (os.path.basename(path), "; ".join(errors)))
    return best


def _extract_docx(path: str) -> Tuple[str, int, str]:
    if docx is None:
        raise MissingBackend(
            "cannot read %s — python-docx is not installed. "
            "Run: python3 -m pip install -r requirements.txt" % os.path.basename(path))
    document = docx.Document(path)
    parts = [p.text for p in document.paragraphs]
    # Fact finds are mostly tables; ignoring them loses the whole document.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    text = "\n".join(parts)
    # Word has no page concept without rendering; approximate for the per-page
    # quality check rather than pretend to know.
    page_count = max(1, len(text) // 1800)
    return text, page_count, "python-docx"


def _extract_txt(path: str) -> Tuple[str, int, str]:
    with open(path, "r") as fh:
        text = fh.read()
    return text, max(1, len(text) // 1800), "plaintext"


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def extract_file(path: str) -> ExtractedDocument:
    """Extract one document from disk."""
    suffix = os.path.splitext(path)[1].lower()
    if suffix in PDF_SUFFIXES:
        raw, pages, backend = _extract_pdf(path)
    elif suffix in WORD_SUFFIXES:
        raw, pages, backend = _extract_docx(path)
    elif suffix in TEXT_SUFFIXES:
        raw, pages, backend = _extract_txt(path)
    else:
        raise ExtractionError(
            "unsupported file type %r for %s (supported: %s)"
            % (suffix, os.path.basename(path), ", ".join(SUPPORTED_SUFFIXES)))

    text = _normalise(raw)
    return ExtractedDocument(
        name=os.path.basename(path),
        text=text,
        page_count=pages,
        source_path=path,
        backend=backend,
    )


def extract_directory(directory: str) -> Tuple[List[ExtractedDocument], List[Dict[str, str]]]:
    """Extract every supported document in a directory (one intake batch).

    Returns the successes and a list of failures. A file that cannot be read is
    reported, never skipped silently — an unreadable document in a compliance
    file is itself a finding.
    """
    documents = []  # type: List[ExtractedDocument]
    failures = []  # type: List[Dict[str, str]]

    if not os.path.isdir(directory):
        raise ExtractionError("not a directory: %s" % directory)

    for entry in sorted(os.listdir(directory)):
        if entry.startswith("."):
            continue
        path = os.path.join(directory, entry)
        if not os.path.isfile(path):
            continue
        if os.path.splitext(entry)[1].lower() not in SUPPORTED_SUFFIXES:
            failures.append({"file": entry, "error": "unsupported file type"})
            continue
        try:
            documents.append(extract_file(path))
        except ExtractionError as exc:
            failures.append({"file": entry, "error": str(exc)})

    return documents, failures


def from_sample_records(records: List[Dict[str, Any]]) -> List[ExtractedDocument]:
    """Wrap sample_documents.json entries as if they had been extracted.

    Keeps the synthetic path and the real path identical downstream, so a fix
    proven against the samples is a fix proven for real files.
    """
    out = []  # type: List[ExtractedDocument]
    for rec in records:
        text = _normalise(rec["text"])
        out.append(ExtractedDocument(
            name=rec["file"],
            text=text,
            page_count=rec.get("pages", 1),
            source_path=None,
            backend="sample",
        ))
    return out


def backend_status() -> Dict[str, bool]:
    """Which optional backends are available — reported in the run header."""
    return {"pypdf": pypdf is not None,
            "pdfplumber": pdfplumber is not None,
            "python-docx": docx is not None}
